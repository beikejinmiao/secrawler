#!/usr/bin/env python
# -*- coding:utf-8 -*-
import re
import os
import json
import shutil
import rarfile
import pdfplumber
from py7zr import pack_7zarchive, unpack_7zarchive
from docx import Document
import pandas as pd
from id_validator import validator
from urllib.parse import urlparse
from utils import tree, tree2list
from utils import reader, traverse
from modules.btbu.util import find_idcard
from paths import DOWNLOADS, DUMP_HOME
from libs.logger import logger

# https://bbs.huaweicloud.com/blogs/180864
# register file format at first.
shutil.register_archive_format('7zip',
                               pack_7zarchive,
                               description='7zip archive')
shutil.register_unpack_format('7zip',
                              ['.7z'],
                              unpack_7zarchive,
                              description='7zip archive')


def _safe_int(item):
    try:
        i = int(item)
    except:
        i = 0
    return i


def load2find(path):
    files = traverse(path)
    stats = tree()   # 统计读取文件成功数量与失败数量
    results = dict()
    for filepath in files:
        filename = os.path.basename(filepath)
        suffix = filename.split('.')[-1].lower()
        candidates = list()
        try:
            if re.match(r'.*\.(txt|csv|xml|json)$', filepath, re.I):
                for line in reader(filepath, raisexp=True):
                    candidates.extend(find_idcard(line))
            elif re.match(r'.*\.doc[x]?$', filepath, re.I):
                logger.info("Load: '%s'" % filepath)
                # docx.opc.exceptions.PackageNotFoundError: Package not found at ''
                doc = Document(filepath)
                for p in doc.paragraphs:
                    candidates.extend(find_idcard(p.text))
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            candidates.extend(find_idcard(cell.text))
            elif re.match(r'.*\.xls[x]?$', filepath, re.I):
                logger.info("Load: '%s'" % filepath)
                # ValueError: Excel file format cannot be determined, you must specify an engine manually.
                xls = pd.read_excel(filepath, sheet_name=None)
                for name, sheet in xls.items():
                    for index, row in sheet.iterrows():
                        for value in row:
                            candidates.extend(find_idcard(str(value)))
            elif re.match(r'.*\.pdf$', filepath, re.I):
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()  # 提取文本
                        candidates.extend(find_idcard(text))
            else:
                continue
            stats[suffix]['success'] = _safe_int(stats[suffix]['success']) + 1
        except Exception as e:
            logger.error(e)
            stats[suffix]['failed'] = _safe_int(stats[suffix]['failed']) + 1
        if len(candidates) > 0:
            results[filename] = list(candidates)
    return results, stats


def unarchive(path):
    stats = tree()  # 统计解压缩文件和读取文件成功数量与失败数量
    files = traverse(path)
    results = tree()
    for filepath in files:
        filename = os.path.basename(filepath)
        suffix = filename.split('.')[-1].lower()
        try:
            dest_dir = ''
            if re.match(r'.*\.(zip|7z|tar|tar\.bz2|tar\.gz|tar\.xz|tbz2|tgz|txz)$', filepath, re.I):
                dest_dir = filepath + '.unpack'
                # shutil.ReadError: xxx.zip is not a zip file
                shutil.unpack_archive(filepath, dest_dir)
                stats[suffix]['success'] = _safe_int(stats[suffix]['success']) + 1
            elif filename.endswith('.rar'):
                # TODO: rarfile.RarCannotExec: Cannot find working tool
                rarfile.UNRAR_TOOL = r"C:\OptSoft\unrar\UnRAR.exe"
                rar = rarfile.RarFile(filepath)
                dest_dir = filepath + '.unpack'
                with rar as rf:
                    rf.extractall(dest_dir)
                stats[suffix]['success'] = _safe_int(stats[suffix]['success']) + 1
            #
            if dest_dir:
                ret, _stats_ = load2find(dest_dir)
                pkg_name = filename
                if len(ret) > 0:
                    results[pkg_name] = ret
            else:
                ret, _stats_ = load2find(filepath)
                pkg_name = ''
                if len(ret) > 0:
                    results[pkg_name][filename] = ret[filename]
                for s in _stats_:
                    stats[suffix]['success'] = _safe_int(stats[suffix]['success']) + _safe_int(_stats_[s]['success'])
                    stats[suffix]['failed'] = _safe_int(stats[suffix]['failed']) + _safe_int(_stats_[s]['failed'])
        except Exception as e:
            logger.error(e)
            stats[suffix]['failed'] = _safe_int(stats[suffix]['failed']) + 1
    logger.info('文件读取成功&失败统计:\n' + json.dumps(dict(stats), indent=4))
    return results


def dump2csv(infos):
    df = pd.DataFrame(tree2list(infos), columns=['archive', 'filename', 'idcard'])
    df = df.assign(idcard=df["idcard"]).explode("idcard").reset_index(drop=True)
    df['is_valid'] = df['idcard'].map(lambda x: 1 if validator.is_valid(x) else 0)
    with open(os.path.join(DUMP_HOME, 'file_urls.json')) as fopen:
        file_urls = json.load(fopen)
    fname2url = dict()
    for file_url in file_urls:
        filename = os.path.basename(urlparse(file_url).path)
        fname2url[filename] = (file_urls[file_url], file_url)
    csv = list()
    for index, row in df.iterrows():
        filename = row['archive'] if row['archive'] else row['filename']
        url, file_url = '', ''
        if filename in fname2url:
            url, file_url = fname2url[filename]
        csv.append(dict(zip(['url', 'file_url', 'archive', 'filename', 'idcard', 'is_valid'],
                            [url, file_url, row['archive'], row['filename'], row['idcard'], row['is_valid']])))
    pd.DataFrame(csv).to_csv(os.path.join(DUMP_HOME, 'file_results.csv'))


if __name__ == '__main__':
    infos = unarchive(DOWNLOADS)
    print('文件内容提取结果:\n', json.dumps(infos))
    dump2csv(infos)




